'''
@author: xusheng
'''

import os
import xlwt
from six.moves import xrange
import argparse
import Levenshtein as lev
import docx

class Sequence(object):
    def __init__(self, seq, **kwargs):
        self._seq = seq
        self._trunc_head = 0
        self._trunc_tail = 0
        
        if 'trunc_head' in kwargs:
            self._trunc_head = int(kwargs['trunc_head'])
            
        if 'trunc_tail' in kwargs:
            self._trunc_tail = int(kwargs['trunc_tail'])
        
        self._trunc_seq_offset = self._trunc_head
        self._trunc_seq_size = len(self._seq) - self._trunc_head - self._trunc_tail

    @property
    def seq(self):
        return self._seq
    
    @property
    def seq_size(self):
        return len(self._seq)
    
    @property
    def trunc_seq(self):
        return self._seq[self._trunc_seq_offset:self._trunc_seq_offset+self._trunc_seq_size]

    @property
    def trunc_seq_offset(self):
        return self._trunc_seq_offset
    
    @trunc_seq_offset.setter
    def trunc_seq_offset(self, value):
        self._trunc_seq_offset = value
    
    @property
    def trunc_seq_size(self):
        return self._trunc_seq_size
    
    @trunc_seq_size.setter
    def trunc_seq_size(self, value):
        self._trunc_seq_size = value


class SSA():
    def __init__(self):
        self._init_ref_sequence()
        self._word_logger = SSARecordWordLogger()
        self._excel_logger = SSARecordExcelLogger()
    
    def _init_ref_sequence(self):
        self._ref_sequence = Sequence(self._load_seq(FLAGS.ref))
    
    def _load_sample_sequence(self, file):
        d_trunc = {'trunc_head': FLAGS.trunc_head, 'trunc_tail': FLAGS.trunc_tail}
        return Sequence(self._load_seq(file), **d_trunc)
    
    def _load_seq(self, file):
        seq = ''
        with open(file, 'r') as f:
            for l in f.readlines():
                seq += l.strip('\n')
        return seq
    
    def _write_seq(self, file, seq):
        with open(file, 'w') as f:
            f.write(seq)
    
    def _check(self, ref_sequence, sample_sequence):
        ret_ld_val = sample_sequence.trunc_seq_size
        
        ref_sequence.trunc_seq_size = sample_sequence.trunc_seq_size
        ref_sequence.trunc_seq_offset = 0
        offset = 0
        
        for i in xrange(ref_sequence.seq_size - sample_sequence.trunc_seq_size):
            ref_sequence.trunc_seq_offset = i
            ld_val = lev.distance(ref_sequence.trunc_seq, sample_sequence.trunc_seq)
            if ld_val < ret_ld_val:
                ret_ld_val = ld_val
                offset = i
        
        ref_sequence.trunc_seq_offset = offset
        
        return ret_ld_val, ref_sequence, sample_sequence
    
    def _match_block(self, ref_seq, sample_seq):
        return lev.matching_blocks(lev.editops(ref_seq, sample_seq), ref_seq, sample_seq)
    
    def process(self):
        filelist = os.listdir(FLAGS.src)
        
        for file in filelist:
            sample_sequence = self._load_sample_sequence(os.path.join(FLAGS.src, file))
            self._write_seq(os.path.join(FLAGS.tar, file), sample_sequence.seq)
            
            idx = file.index('-')
            patient_no = file[0:idx]
            
            ld_val, self._ref_sequence, sample_sequence = self._check(self._ref_sequence, sample_sequence)
            
            if ld_val == 0:
                d_record = {
                    'reaction': FLAGS.reaction,
                    'patient_no': patient_no,
                    'ld': ld_val
                }
            else:
                mb = self._match_block(self._ref_sequence.trunc_seq, sample_sequence.trunc_seq)
                d_record = {
                    'reaction': FLAGS.reaction,
                    'patient_no': patient_no,
                    'ld': ld_val,
                    'ref_sequence': self._ref_sequence,
                    'sample_sequence': sample_sequence,
                    'match_block': mb
                }
                self._word_logger.log(SSARecord(**d_record))
        
            self._excel_logger.log(SSARecord(**d_record))
            
        self._word_logger.flush()
        self._excel_logger.flush()
            

class SSARecord(object):
    def __init__(self, **kwargs):
        self._items = {
            'reaction': '',
            'patient_no': '' ,
            'ld': 0,
            'sample_sequence': None,
            'ref_sequence': None,
            'match_block': None
        }
        
        for kwarg in kwargs:
            if kwarg not in self._items:
                raise TypeError('Keyword argument not understood:', kwarg)
            self._items[kwarg] = kwargs[kwarg]
        
    @property
    def items(self):
        return self._items

class Logger(object):
    def __init__(self):
        self._init_logger()
    
    def _init_logger(self):
        pass
    
    def log(self, obj):
        if isinstance(obj, str):
            print(obj)
    
    def flush(self):
        pass
    
class SSARecordWordLogger(Logger):
    def _init_logger(self):
        self._document = docx.Document()
        self._init_style()
        self._init_reference_paragraph()
    
    def _init_style(self):
        self._header_style = self._document.styles.add_style('Customized Header', 1)
        self._header_style.font.size = docx.shared.Pt(24)
        self._header_style.font.bold = True
        self._header_style.font.name = 'Courier New'
        self._header_style.font.color.rgb = docx.shared.RGBColor(0x0, 0x0, 0x0)

        self._normal_style = self._document.styles.add_style('Customized Normal', 2)
        self._normal_style.font.size = docx.shared.Pt(11)
        self._normal_style.font.name = 'Courier New'
        self._normal_style.font.color.rgb = docx.shared.RGBColor(0x0, 0x0, 0x0)

        self._matchblock_style = self._document.styles.add_style('Customized Match Block', 2)
        self._matchblock_style.font.size = docx.shared.Pt(11)
        self._matchblock_style.font.bold = True
        self._matchblock_style.font.name = 'Courier New'
        self._matchblock_style.font.color.rgb = docx.shared.RGBColor(0x0, 0x0, 0x0)
        
        self._highlighted_style = self._document.styles.add_style('Customized Highlighted', 2)
        self._highlighted_style.font.size = docx.shared.Pt(11)
        self._highlighted_style.font.bold = True
        self._highlighted_style.font.name = 'Courier New'
        self._highlighted_style.font.color.rgb = docx.shared.RGBColor(0x0, 0x0, 0xff)
    
    def _init_reference_paragraph(self):
        self._document.add_paragraph('Report', style='Customized Header')
        self._document.add_paragraph().add_run('Reaction: %s' % (FLAGS.reaction), style='Customized Normal')
        self._document.add_paragraph().add_run('---------------------------------------------------', style='Customized Normal')
        
    def log(self, obj):
        if isinstance(obj, SSARecord):
            self._document.add_paragraph().add_run('#[%s]' % (obj.items['patient_no']), style='Customized Normal')
            
            sample_sequence = obj.items['sample_sequence']
            ref_sequence = obj.items['ref_sequence']
            
            # init two paragraphs: ref, sample
            p_ref = self._document.add_paragraph()
            p_ref.add_run('REF   : ', style='Customized Normal')
            
            p_sample = self._document.add_paragraph()
            p_sample.add_run('SAMPLE: ', style='Customized Normal')
            
            # log block before match_block
            offset = ref_sequence.trunc_seq_offset - sample_sequence.trunc_seq_offset
            if offset > 0:
                p_ref.add_run(ref_sequence.seq[0:ref_sequence.trunc_seq_offset], style='Customized Normal')
                
                blank = ''.join([" " for x in xrange(offset)])
                p_sample.add_run(blank, style='Customized Normal')
                p_sample.add_run(sample_sequence.seq[0:sample_sequence.trunc_seq_offset], style='Customized Normal')
            elif offset < 0:
                blank = ''.join([" " for x in xrange(-offset)])
                p_ref.add_run(blank, style='Customized Normal')
                p_ref.add_run(ref_sequence.seq[0:ref_sequence.trunc_seq_offset], style='Customized Normal')
                
                p_sample.add_run(sample_sequence.seq[0:sample_sequence.trunc_seq_offset], style='Customized Normal')
            
            # log match block
            mb = obj.items['match_block']
            ref_cursor = 0
            sample_cursor = 0
            
            for x in mb:
                # not match
                ref_not_match = ref_sequence.trunc_seq[ref_cursor:x[0]]
                if ref_not_match != '':
                    p_ref.add_run(ref_not_match, style='Customized Highlighted')

                sample_not_match = sample_sequence.trunc_seq[sample_cursor:x[1]]
                if sample_not_match != '':
                    p_sample.add_run(sample_not_match, style='Customized Highlighted')

                # match
                ref_match = ref_sequence.trunc_seq[x[0]:x[0]+x[2]]
                if ref_match != '':
                    p_ref.add_run(ref_match, style='Customized Match Block')

                sample_match = sample_sequence.trunc_seq[x[1]:x[1]+x[2]]
                if sample_match != '':
                    p_sample.add_run(sample_match, style='Customized Match Block')
                
                ref_cursor = x[0]+x[2]
                sample_cursor = x[1]+x[2]

            # log rest info
            p_ref.add_run(ref_sequence.seq[ref_sequence.trunc_seq_offset+ref_sequence.trunc_seq_size:], style='Customized Normal')
            p_sample.add_run(sample_sequence.seq[sample_sequence.trunc_seq_offset+sample_sequence.trunc_seq_size:], style='Customized Normal')

            self._document.add_paragraph().add_run('---------------------------------------------------', style='Customized Normal')
        else:
            pass
    
    def flush(self):
        self._document.save(FLAGS.word_log)
    
class SSARecordExcelLogger(Logger):
    def _init_logger(self):
        self._workbook = xlwt.Workbook(encoding='utf-8')
        self._sheet = self._workbook.add_sheet('report', cell_overwrite_ok=True)
        self._sheet.col(0).width = 3000
        self._sheet.col(1).width = 4000
        self._sheet.col(2).width = 1000
        self._sheet.col(3).width = 10000
        self._sheet.col(4).width = 10000
        
        self._cursor = 0
        self._init_header(self._sheet)
    
    def _header_info(self):
        headers = [['Reaction', 'Patient No', 'LD', 'Truncated Sample', 'Truncated Reference'],]
        return headers, (len(headers), len(headers[0]))
    
    def _init_header(self, sheet):
        header_style = xlwt.XFStyle()
        header_style.font.bold = True
        
        header, shape = self._header_info()
        for row in xrange(shape[0]):
            for col in xrange(shape[1]):
                sheet.write(self._cursor, col, label=header[row][col], style=header_style)
            self._cursor += 1
    
    def log(self, obj):
        if isinstance(obj, SSARecord):
            self._sheet.write(self._cursor, 0, label=obj.items['reaction'])
            self._sheet.write(self._cursor, 1, label=obj.items['patient_no'])
            self._sheet.write(self._cursor, 2, label=obj.items['ld'])
            
            sample_sequence = obj.items['sample_sequence']
            if sample_sequence is not None:
                self._sheet.write(self._cursor, 3, label=sample_sequence.trunc_seq)
                
            ref_sequence = obj.items['ref_sequence']
            if ref_sequence is not None:
                self._sheet.write(self._cursor, 4, label=ref_sequence.trunc_seq)
            
            self._cursor += 1
        else:
            pass
    
    def flush(self):
        self._workbook.save(FLAGS.excel_log)


if __name__=='__main__':
    global FLAGS
    parser = argparse.ArgumentParser(description='Sanger Sequencing Analysis Toolkit Helper:')
    parser.add_argument('--ref', dest='ref', type=str, default=os.path.join('.', 'ref', '2.2F.txt'), help='reference file')
    parser.add_argument('--src', dest='src', type=str, default=os.path.join('.', 'src'), help='source file path')
    parser.add_argument('--tar', dest='tar', type=str, default=os.path.join('.', 'tar'), help='target file path')
    parser.add_argument('--excel_log', dest='excel_log', type=str, default=os.path.join('.', 'log', 'report.xls'), help='log file with XLS format')
    parser.add_argument('--word_log', dest='word_log', type=str, default=os.path.join('.', 'log', 'report.docx'), help='log file with DOCX format')
    parser.add_argument('--trunc_head', dest='trunc_head', type=int, default=30, help='numbers for head truncation')
    parser.add_argument('--trunc_tail', dest='trunc_tail', type=int, default=10, help='numbers for tail truncation')
    parser.add_argument('--reaction', dest='reaction', type=str, default='UNDEFINED', help='reaction category')
    
    FLAGS = parser.parse_args()
  
    ssa = SSA()
    ssa.process()
